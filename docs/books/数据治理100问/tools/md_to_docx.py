#!/usr/bin/env python3
from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BOOK_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_OUTPUT_DIR = BOOK_ROOT / "发布稿"
BODY_FONT = "Arial Unicode MS"
HEADING_FONT = "Arial Unicode MS"
BLUE = "2E5D7B"
DARK_BLUE = "203E52"
MUTED = "687580"
LIGHT_BLUE = "EEF4F7"
TABLE_HEADER = "E8EEF2"
BORDER = "B8C6CF"
CONTENT_WIDTH_DXA = 9360

CHINESE_NUMERALS = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
    11: "十一",
    12: "十二",
}


def find_executable(name: str, fallbacks: list[str] | None = None) -> str:
    resolved = shutil.which(name)
    if resolved:
        return resolved
    for candidate in fallbacks or []:
        if Path(candidate).is_file():
            return candidate
    raise RuntimeError(f"未找到命令：{name}")


def run_checked(command: list[str], env: dict[str, str] | None = None) -> None:
    subprocess.run(command, check=True, env=env)


def set_run_font(
    run,
    font: str,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def table_widths(table) -> list[int]:
    count = len(table.columns)
    if count == 1:
        return [CONTENT_WIDTH_DXA]
    if count == 2:
        first_max = max(len(row.cells[0].text.strip()) for row in table.rows)
        first = 2160 if first_max <= 16 else 3200
        return [first, CONTENT_WIDTH_DXA - first]
    if count == 3:
        return [1800, 3780, 3780]
    base = CONTENT_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, HEADING_FONT, size=9, color=MUTED)


def set_footer(paragraph) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    left = paragraph.add_run("—  ")
    set_run_font(left, HEADING_FONT, size=9, color=MUTED)
    add_page_field(paragraph)
    right = paragraph.add_run("  —")
    set_run_font(right, HEADING_FONT, size=9, color=MUTED)


def shade_paragraph(paragraph, fill: str, border_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


NODE_PATTERN = r'([A-Za-z0-9_]+)(?:\["([^"]*)"\])?'


def parse_node(token: str) -> tuple[str, str | None]:
    match = re.fullmatch(NODE_PATTERN, token.strip())
    if not match:
        raise ValueError(f"无法识别 Mermaid 节点：{token}")
    return match.group(1), match.group(2)


def mermaid_to_dot(source: str) -> str:
    lines = [line.strip() for line in source.splitlines() if line.strip()]
    if not lines:
        raise ValueError("Mermaid 图为空")
    head = re.fullmatch(r"flowchart\s+(LR|RL|TB|BT)", lines[0])
    if not head:
        raise ValueError("当前仅支持 Mermaid flowchart LR/RL/TB/BT")
    direction = head.group(1)
    nodes: dict[str, str] = {}
    edges: list[tuple[str, str, str, str | None]] = []

    patterns = (
        (re.compile(r"^(.+?)\s+-\.\s*(.*?)\s*\.->\s*(.+?)$"), "dotted"),
        (re.compile(r"^(.+?)\s*<-->\s*(.+?)$"), "both"),
        (re.compile(r"^(.+?)\s*-->\s*(.+?)$"), "solid"),
    )

    for line in lines[1:]:
        matched = False
        for pattern, kind in patterns:
            match = pattern.fullmatch(line)
            if not match:
                continue
            groups = match.groups()
            if kind == "dotted":
                left_token = groups[0].strip()
                label = groups[1].strip() or None
                right_token = groups[2].strip()
            else:
                left_token = groups[0].strip()
                label = None
                right_token = groups[1].strip()
            left_id, left_label = parse_node(left_token)
            right_id, right_label = parse_node(right_token)
            nodes.setdefault(left_id, (left_label or left_id).replace(r"\n", "\n"))
            nodes.setdefault(right_id, (right_label or right_id).replace(r"\n", "\n"))
            edges.append((left_id, right_id, kind, label))
            matched = True
            break
        if not matched:
            raise ValueError(f"当前转换器无法识别 Mermaid 语句：{line}")

    rankdir = {"LR": "LR", "RL": "RL", "TB": "TB", "BT": "BT"}[direction]
    output = [
        "digraph G {",
        f"  rankdir={rankdir};",
        '  graph [bgcolor="white", pad="0.2", nodesep="0.35", ranksep="0.55"];',
        '  node [shape=box, style="rounded,filled", fillcolor="#EAF2F6", color="#2E5D7B", penwidth=1.6, fontname="Arial Unicode MS", fontsize=14, fontcolor="#203E52", margin="0.16,0.10"];',
        '  edge [color="#667F8E", penwidth=1.5, arrowsize=0.75, fontname="Arial Unicode MS", fontsize=11, fontcolor="#687580"];',
    ]
    for node_id, label in nodes.items():
        escaped = label.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n")
        output.append(f'  {node_id} [label="{escaped}"];')
    for left, right, kind, label in edges:
        attributes = []
        if kind == "dotted":
            attributes.append("style=dashed")
        if kind == "both":
            attributes.append("dir=both")
        if label:
            escaped_label = label.replace('"', '\\"')
            attributes.append(f'label="{escaped_label}"')
        suffix = f" [{', '.join(attributes)}]" if attributes else ""
        output.append(f"  {left} -> {right}{suffix};")
    output.append("}")
    return "\n".join(output)


def render_mermaid(source: str, destination: Path) -> None:
    dot = find_executable("dot", ["/opt/homebrew/bin/dot", "/usr/local/bin/dot"])
    dot_source = destination.with_suffix(".dot")
    dot_source.write_text(mermaid_to_dot(source), encoding="utf-8")
    env = os.environ.copy()
    env.setdefault("LANG", "zh_CN.UTF-8")
    run_checked([dot, "-Tpng", "-Gdpi=220", str(dot_source), "-o", str(destination)], env=env)


def rewrite_local_docx_links(content: str) -> str:
    pattern = re.compile(r"\]\(([^)]+?\.md)(#[^)]+)?\)")

    def replace(match: re.Match[str]) -> str:
        target = Path(match.group(1)).name
        return f"]({Path(target).with_suffix('.docx').as_posix()}{match.group(2) or ''})"

    return pattern.sub(replace, content)


def prepare_markdown(source: Path, temp_dir: Path) -> Path:
    content = rewrite_local_docx_links(source.read_text(encoding="utf-8"))
    counter = 0

    def replace_mermaid(match: re.Match[str]) -> str:
        nonlocal counter
        counter += 1
        image_path = temp_dir / f"diagram-{counter}.png"
        render_mermaid(match.group(1), image_path)
        return f"![流程图]({image_path.as_posix()})"

    content = re.sub(r"```mermaid\n(.*?)\n```", replace_mermaid, content, flags=re.S)
    prepared = temp_dir / "article.md"
    prepared.write_text(content, encoding="utf-8")
    return prepared


def chapter_label(source: Path) -> str:
    match = re.match(r"(\d+)-(.+?)篇$", source.parent.name)
    if not match:
        return "书稿大纲"
    number = int(match.group(1))
    name = match.group(2)
    return f"第{CHINESE_NUMERALS.get(number, str(number))}篇 {name}"


def style_document(base_docx: Path, source: Path, output: Path) -> None:
    doc = Document(base_docx)
    if not doc.paragraphs:
        raise RuntimeError(f"转换后文档为空：{source}")
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.86)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("Block Text", "Caption"):
        if style_name in styles:
            styles[style_name].font.name = BODY_FONT
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    title = doc.paragraphs[0]
    title.style = styles["Normal"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    is_outline = source.name == "大纲.md"
    for run in title.runs:
        set_run_font(run, HEADING_FONT, size=24 if not is_outline else 26, color=DARK_BLUE, bold=True)

    label = chapter_label(source)
    kicker = title.insert_paragraph_before()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(8)
    kicker_run = kicker.add_run(f"数据治理100问  ·  {label}")
    set_run_font(kicker_run, HEADING_FONT, size=9.5, color=BLUE, bold=True)

    for paragraph in doc.paragraphs:
        if paragraph is title or paragraph is kicker:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower() in {"source code", "verbatim", "code block"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, size=9.5, color="333333")
            continue
        if style_name.startswith("Heading"):
            size, color = heading_tokens.get(style_name, (11, BLUE, 0, 0))[:2]
            for run in paragraph.runs:
                set_run_font(run, HEADING_FONT, size=size, color=color, bold=True)
        else:
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, size=11, color="000000")
            p_pr = paragraph._p.pPr
            if p_pr is not None and p_pr.numPr is not None:
                paragraph.paragraph_format.left_indent = Inches(0.38)
                paragraph.paragraph_format.first_line_indent = Inches(-0.19)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.208

    for paragraph in doc.paragraphs:
        if not paragraph.style or paragraph.style.name != "Block Text":
            continue
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.12)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(12)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.keep_together = True
        shade_paragraph(paragraph, LIGHT_BLUE, BLUE)
        for run in paragraph.runs:
            set_run_font(run, BODY_FONT, size=11.5, color=DARK_BLUE, bold=True)

    for table in doc.tables:
        set_table_geometry(table, table_widths(table))
        set_table_borders(table)
        set_repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if row_index == 0:
                    set_cell_shading(cell, TABLE_HEADER)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.15
                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            HEADING_FONT if row_index == 0 else BODY_FONT,
                            size=10.2,
                            color=DARK_BLUE if row_index == 0 else "000000",
                            bold=(row_index == 0),
                        )

    for shape in doc.inline_shapes:
        width_in = shape.width / 914400
        height_in = shape.height / 914400
        ratio = height_in / width_in if width_in else 0.3
        target_width = 6.15
        target_height = target_width * ratio
        if target_height > 2.75:
            target_height = 2.75
            target_width = target_height / ratio
        shape.width = Inches(target_width)
        shape.height = Inches(target_height)

    for paragraph in doc.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_together = True

    running_label = f"数据治理100问  ·  {label}"
    header = section.header.paragraphs[0]
    header.text = running_label
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, HEADING_FONT, size=8.5, color=MUTED)

    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    set_footer(section.footer.paragraphs[0])
    set_footer(section.first_page_footer.paragraphs[0])

    properties = doc.core_properties
    properties.title = title.text
    properties.subject = "《数据治理100问》发布稿"
    properties.keywords = "数据治理, 业务增值, 数据治理100问"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def convert(source: Path, output_dir: Path) -> Path:
    source = source.resolve()
    if not source.is_file() or source.suffix.lower() != ".md":
        raise FileNotFoundError(f"不是有效的 Markdown 文件：{source}")
    pandoc = find_executable("pandoc", ["/opt/homebrew/bin/pandoc", "/usr/local/bin/pandoc"])
    output = output_dir / f"{source.stem}.docx"
    with tempfile.TemporaryDirectory(prefix="data_governance_docx_") as temp:
        temp_dir = Path(temp)
        prepared = prepare_markdown(source, temp_dir)
        base_docx = temp_dir / "base.docx"
        run_checked([pandoc, str(prepared), "--from=gfm", "--to=docx", "--output", str(base_docx)])
        style_document(base_docx, source, output)
    return output


def all_sources() -> list[Path]:
    articles = sorted(path for path in BOOK_ROOT.glob("[0-9][0-9]-*篇/*.md") if path.is_file())
    outline = BOOK_ROOT / "大纲.md"
    return ([outline] if outline.is_file() else []) + articles


def resolve_source(value: str) -> Path:
    candidate = Path(value)
    if not candidate.is_absolute():
        candidate = BOOK_ROOT / candidate
    return candidate


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="将《数据治理100问》Markdown 源稿转换为统一版式 DOCX")
    parser.add_argument("sources", nargs="*", help="要转换的 Markdown 文件，可使用相对书籍根目录的路径")
    parser.add_argument("--all", action="store_true", help="转换大纲和所有篇章目录下的 Markdown 文件")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR, help="DOCX 输出目录")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.all and args.sources:
        print("错误：--all 不能和指定文件同时使用。", file=sys.stderr)
        return 2
    sources = all_sources() if args.all else [resolve_source(value) for value in args.sources]
    if not sources:
        print("错误：请指定 Markdown 文件，或使用 --all。", file=sys.stderr)
        return 2
    output_dir = args.output_dir
    if not output_dir.is_absolute():
        output_dir = (BOOK_ROOT / output_dir).resolve()
    for source in sources:
        result = convert(source, output_dir)
        print(f"已生成：{result}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
